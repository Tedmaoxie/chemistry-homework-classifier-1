"""
Web 界面：上传题目，生成并下载分层PDF报告
"""

import io
import tempfile
import os
from pathlib import Path
from typing import List, Tuple

from flask import Flask, request, render_template, send_file, abort, jsonify
from flask_cors import CORS

import sys
from pathlib import Path

# Ensure both local src and embedded classifier package are importable
_BASE_DIR = Path(__file__).resolve().parents[2]
_EXTRA_PATHS = [
    _BASE_DIR / "src",
    _BASE_DIR / "chemistry-homework-classifier" / "src",
]
for _p in _EXTRA_PATHS:
    _p_str = str(_p)
    if _p.exists() and _p_str not in sys.path:
        sys.path.insert(0, _p_str)

try:
    from src.difficulty_assessor.classifier import ThreeTierClassifier
except ImportError:
    # Fallback: dynamically load classifier modules under 'src.' namespace
    import importlib.util

    _CC_SRC = _BASE_DIR / "chemistry-homework-classifier" / "src"

    def _load_module(module_name: str, file_path: Path, package_dir: Path = None):
        if package_dir is not None:
            spec = importlib.util.spec_from_file_location(
                module_name, str(file_path), submodule_search_locations=[str(package_dir)]
            )
        else:
            spec = importlib.util.spec_from_file_location(module_name, str(file_path))
        module = importlib.util.module_from_spec(spec)
        sys.modules[module_name] = module
        spec.loader.exec_module(module)
        return module

    # Ensure dependent modules are registered for relative imports
    # Load data package
    # Register packages properly
    _load_module("src.data", _CC_SRC / "data" / "__init__.py", _CC_SRC / "data")
    _load_module("src.data.chemistry_vocab", _CC_SRC / "data" / "chemistry_vocab" / "__init__.py", _CC_SRC / "data" / "chemistry_vocab")
    _load_module("src.difficulty_assessor", _CC_SRC / "difficulty_assessor" / "__init__.py", _CC_SRC / "difficulty_assessor")
    # Load modules
    _load_module("src.difficulty_assessor.assessor", _CC_SRC / "difficulty_assessor" / "assessor.py")
    _classifier_mod = _load_module("src.difficulty_assessor.classifier", _CC_SRC / "difficulty_assessor" / "classifier.py")
    ThreeTierClassifier = _classifier_mod.ThreeTierClassifier
from output_generator.converter import convert_text_bytes, ConversionOptions, decode_text
from web_interface.storage import save_file, classify_and_validate, ensure_storage_base


def parse_blocks(text: str) -> List[Tuple[str, str]]:
    """解析文本为 (题目, 答案) 列表。空行分隔题块。支持“题目：”与“答案：”标记。"""
    content = text or ""
    blocks: List[str] = []
    current: List[str] = []
    for line in content.splitlines():
        if line.strip():
            current.append(line)
        else:
            if current:
                blocks.append("\n".join(current))
                current = []
    if current:
        blocks.append("\n".join(current))

    pairs: List[Tuple[str, str]] = []
    for b in blocks:
        q = []
        a = []
        for ln in b.splitlines():
            s = ln.strip()
            if s.startswith("答案：") or s.startswith("答案:"):
                a.append(s.split("：", 1)[-1] if "：" in s else s.split(":", 1)[-1])
            elif s.startswith("题目：") or s.startswith("题目:"):
                q.append(s.split("：", 1)[-1] if "：" in s else s.split(":", 1)[-1])
            else:
                q.append(ln)
        question_text = "\n".join(q).strip()
        answer_text = "\n".join(a).strip()
        if question_text:
            pairs.append((question_text, answer_text))
    return pairs


def create_app() -> Flask:
    app = Flask(__name__, template_folder=str(Path(__file__).parent / "templates"))
    CORS(app)
    # 允许典型 PDF/DOCX 体量（25MB）
    app.config['MAX_CONTENT_LENGTH'] = 25 * 1024 * 1024
    upload_base = _BASE_DIR / "data" / "uploads"
    ensure_storage_base(upload_base)

    classifier = ThreeTierClassifier()

    @app.get("/")
    def index():
        return render_template("index.html")

    @app.post("/generate")
    def generate():
        # 获取文本或文件
        text = request.form.get("questions_text", "")
        file = request.files.get("file")

        if file and file.filename:
            data = file.read()
            fname = file.filename
            ext = Path(fname).suffix.lower()
            if ext not in {".txt", ".pdf", ".docx"}:
                abort(400, description="仅支持TXT、PDF和DOCX格式")
            # 分类存储原文件
            try:
                saved_path = save_file(data, fname, upload_base)
            except ValueError as e:
                abort(400, description=str(e))
            # 提取文本（支持 PDF/DOCX）
            if ext == ".pdf":
                try:
                    import pdfplumber  # type: ignore
                    import io as _io
                    with pdfplumber.open(_io.BytesIO(data)) as pdf:
                        pages_text = [p.extract_text() or "" for p in pdf.pages]
                    text = "\n\n".join(pages_text)
                except Exception as e:
                    abort(400, description=f"无法解析PDF文本：{e}")
            elif ext == ".docx":
                try:
                    import docx  # type: ignore
                    import io as _io
                    doc = docx.Document(_io.BytesIO(data))
                    text = "\n".join(par.text for par in doc.paragraphs)
                except Exception as e:
                    abort(400, description=f"无法解析DOCX文本：{e}")
            else:  # .txt
                try:
                    text = decode_text(data)
                except Exception:
                    abort(400, description="无法读取上传文件，请使用UTF-8或GBK编码的纯文本文件。")

        if not text.strip():
            abort(400, description="请提供题目文本或上传文件。")

        pairs = parse_blocks(text)
        if not pairs:
            abort(400, description="未解析到题目，请检查格式。题目块之间需空行分隔。")

        # 延迟导入输出模块，保证更友好的错误提示
        try:
            from output_generator import ClassifiedQuestion, ClassifiedPDFGenerator, PDFReportStyle
        except Exception as e:
            abort(500, description=f"输出模块导入失败，请先安装依赖：python -m pip install -r requirements.txt。错误：{e}")

        items: List[ClassifiedQuestion] = []
        for q, a in pairs:
            res = classifier.classify_question(q, a or None)
            items.append(ClassifiedQuestion(question_text=q, answer_text=a or None, classification=res))

        # 生成PDF到内存并返回下载（按综合难度评分降序）
        style = PDFReportStyle(title="高中化学分层作业报告", author="chemistry-homework-classifier")
        generator = ClassifiedPDFGenerator(style=style)

        # Windows 兼容：避免 NamedTemporaryFile 的句柄占用导致写入报错
        fd, tmp_path = tempfile.mkstemp(suffix=".pdf")
        os.close(fd)
        try:
            generator.generate_sorted_report(tmp_path, items, metadata={
                "subtitle": "Web生成的分层与难度分析",
                "sort_note": "排序依据：综合难度评分降序",
            })
            pdf_bytes = Path(tmp_path).read_bytes()
        finally:
            try:
                Path(tmp_path).unlink()
            except Exception:
                pass

        buf = io.BytesIO(pdf_bytes)
        buf.seek(0)
        return send_file(buf, mimetype="application/pdf", as_attachment=True, download_name="classified_report.pdf")

    @app.post("/convert")
    def convert():
        # 目标格式
        target = (request.form.get("target_format") or request.args.get("target_format") or "pdf").lower()
        if target not in ("pdf", "docx"):
            abort(400, description="不支持的目标格式，仅支持 pdf 或 docx")

        # 获取文件或文本
        file = request.files.get("file")
        if not file or not file.filename:
            # 兼容直接文本
            text = request.form.get("text_content", "")
            if not text.strip():
                abort(400, description="请上传文件或提供文本内容。")
            data = text.encode("utf-8")
        else:
            data = file.read()
            fname = file.filename
            ext = Path(fname).suffix.lower()
            if ext not in {".txt", ".pdf", ".docx"}:
                abort(400, description="仅支持TXT、PDF和DOCX格式")
            # 提取或解码文本内容
            if ext == ".pdf":
                try:
                    import pdfplumber  # type: ignore
                    import io as _io
                    with pdfplumber.open(_io.BytesIO(data)) as pdf:
                        pages_text = [p.extract_text() or "" for p in pdf.pages]
                    data = "\n\n".join(pages_text).encode("utf-8")
                except Exception as e:
                    abort(400, description=f"无法解析PDF文本：{e}")
            elif ext == ".docx":
                try:
                    import docx  # type: ignore
                    import io as _io
                    doc = docx.Document(_io.BytesIO(data))
                    text = "\n".join(par.text for par in doc.paragraphs)
                    data = text.encode("utf-8")
                except Exception as e:
                    abort(400, description=f"无法解析DOCX文本：{e}")
            else:  # .txt
                try:
                    text = decode_text(data)
                    data = text.encode("utf-8")
                except Exception:
                    abort(400, description="无法读取上传文件，请使用UTF-8或GBK编码的纯文本文件。")

        try:
            options = ConversionOptions(target_format=target, title="上传文本转换文档")
            suffix = ".pdf" if target == "pdf" else ".docx"
            fd, tmp_path = tempfile.mkstemp(suffix=suffix)
            os.close(fd)
            try:
                convert_text_bytes(data, tmp_path, options)
                out_bytes = Path(tmp_path).read_bytes()
            finally:
                try:
                    Path(tmp_path).unlink()
                except Exception:
                    pass
        except TimeoutError:
            abort(503, description="转换超时，请减小文件或稍后重试。")
        except Exception as e:
            abort(500, description=f"转换失败：{e}")

        mime = "application/pdf" if target == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        download_name = "converted.pdf" if target == "pdf" else "converted.docx"
        buf = io.BytesIO(out_bytes)
        buf.seek(0)
        return send_file(buf, mimetype=mime, as_attachment=True, download_name=download_name)

    @app.post("/upload")
    def upload():
        file = request.files.get("file")
        if not file or not file.filename:
            abort(400, description="请上传文件")
        data = file.read()
        try:
            ftype = classify_and_validate(data, file.filename)
        except ValueError as e:
            abort(400, description=str(e))
        saved_path = save_file(data, file.filename, upload_base)
        return jsonify({
            "status": "ok",
            "file_type": ftype,
            "stored_path": str(saved_path),
            "message": "上传成功"
        })

    return app


app = create_app()

if __name__ == "__main__":
    # 开发模式运行
    app.run(host="127.0.0.1", port=5000, debug=True)