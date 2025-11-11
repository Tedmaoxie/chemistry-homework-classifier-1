"""
文档转换器：将上传的文本内容转换为标准 PDF / DOCX 文档

支持：
- 自动编码识别（chardet）
- 段落、标题、列表解析
- PDF（reportlab）与 DOCX（python-docx）生成
"""

import logging
import time
from dataclasses import dataclass
from typing import Optional, List, Dict, Any
from pathlib import Path
import re

try:
    import chardet  # type: ignore
except Exception:
    chardet = None

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    ListFlowable,
    ListItem,
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

logger = logging.getLogger(__name__)


@dataclass
class ConversionOptions:
    target_format: str  # 'pdf' or 'docx'
    title: str = "转换文档"
    author: str = "chemistry-homework-classifier"
    font_name: str = "STSong-Light"
    font_path: Optional[str] = None
    page_size: Any = A4
    timeout_seconds: int = 30


def detect_encoding(data: bytes) -> str:
    """检测文本编码，返回用于解码的编码名称。"""
    if chardet:
        try:
            info = chardet.detect(data)
            enc = info.get("encoding") or "utf-8"
            return enc
        except Exception:
            return "utf-8"
    return "utf-8"


def decode_text(data: bytes) -> str:
    enc = detect_encoding(data)
    try:
        return data.decode(enc)
    except Exception:
        for fallback in ("utf-8", "gbk", "latin-1"):
            try:
                return data.decode(fallback)
            except Exception:
                continue
    return data.decode("utf-8", errors="ignore")


def parse_structure(text: str) -> List[Dict[str, Any]]:
    """将纯文本解析为结构化块：heading/paragraph/list。"""
    blocks: List[Dict[str, Any]] = []
    lines = text.splitlines()
    current_par: List[str] = []
    list_buffer: List[str] = []
    list_type: Optional[str] = None  # 'ul' or 'ol'

    def flush_par():
        nonlocal current_par
        if current_par:
            blocks.append({"type": "paragraph", "text": "\n".join(current_par)})
            current_par = []

    def flush_list():
        nonlocal list_buffer, list_type
        if list_buffer:
            blocks.append({"type": "list", "items": list_buffer[:], "ordered": (list_type == "ol")})
            list_buffer = []
            list_type = None

    for ln in lines:
        s = ln.strip()
        # Markdown 风格标题
        if s.startswith("### "):
            flush_par(); flush_list()
            blocks.append({"type": "heading", "level": 3, "text": s[4:].strip()})
            continue
        elif s.startswith("## "):
            flush_par(); flush_list()
            blocks.append({"type": "heading", "level": 2, "text": s[3:].strip()})
            continue
        elif s.startswith("# "):
            flush_par(); flush_list()
            blocks.append({"type": "heading", "level": 1, "text": s[2:].strip()})
            continue

        # 列表项
        if re.match(r"^([*\-•])\s+", s):
            flush_par()
            item = re.sub(r"^([*\-•])\s+", "", s)
            if list_type not in (None, "ul"):
                flush_list()
            list_type = "ul"
            list_buffer.append(item)
            continue

        if re.match(r"^\d+[\.)]\s+", s):
            flush_par()
            item = re.sub(r"^\d+[\.)]\s+", "", s)
            if list_type not in (None, "ol"):
                flush_list()
            list_type = "ol"
            list_buffer.append(item)
            continue

        # 空行分段
        if not s:
            flush_list(); flush_par()
            continue

        # 普通文本归为段落
        current_par.append(ln)

    flush_list(); flush_par()
    return blocks


def _register_fonts(options: ConversionOptions) -> str:
    try:
        if options.font_path:
            pdfmetrics.registerFont(TTFont(options.font_name, options.font_path))
            return options.font_name
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        return "STSong-Light"
    except Exception:
        return "Helvetica"


def convert_to_pdf(text: str, output_path: str, options: ConversionOptions) -> str:
    start = time.time()
    font_name = _register_fonts(options)
    base = getSampleStyleSheet()
    title = ParagraphStyle(name="Title", parent=base["Title"], fontName=font_name, fontSize=20, leading=24, textColor=colors.HexColor("#1f6feb"), alignment=1)
    h1 = ParagraphStyle(name="Heading1", parent=base["Heading1"], fontName=font_name, fontSize=16, leading=20)
    h2 = ParagraphStyle(name="Heading2", parent=base["Heading2"], fontName=font_name, fontSize=14, leading=18)
    h3 = ParagraphStyle(name="Heading3", parent=base["Heading3"], fontName=font_name, fontSize=12, leading=16)
    body = ParagraphStyle(name="Body", parent=base["BodyText"], fontName=font_name, fontSize=11, leading=16)

    doc = SimpleDocTemplate(
        output_path,
        pagesize=options.page_size,
        leftMargin=36,
        rightMargin=36,
        topMargin=48,
        bottomMargin=48,
        title=options.title,
        author=options.author,
    )

    blocks = parse_structure(text)
    story: List[Any] = [Paragraph(options.title, title), Spacer(1, 12)]
    for b in blocks:
        if time.time() - start > options.timeout_seconds:
            raise TimeoutError("PDF 转换超时")
        if b["type"] == "heading":
            lvl = b.get("level", 2)
            style = h1 if lvl == 1 else h2 if lvl == 2 else h3
            story.append(Paragraph(b["text"], style))
            story.append(Spacer(1, 6))
        elif b["type"] == "paragraph":
            story.append(Paragraph(b["text"].replace("\n", "<br/>"), body))
            story.append(Spacer(1, 6))
        elif b["type"] == "list":
            bulletType = "1" if b.get("ordered") else "bullet"
            items = [ListItem(Paragraph(it, body)) for it in b.get("items", [])]
            story.append(ListFlowable(items, bulletType=bulletType))
            story.append(Spacer(1, 6))

    doc.build(story)
    return output_path


def convert_to_docx(text: str, output_path: str, options: ConversionOptions) -> str:
    start = time.time()
    try:
        import docx  # type: ignore
        from docx.shared import Pt
        from docx.oxml.ns import qn
    except Exception as e:
        raise RuntimeError(f"DOCX 转换不可用，请安装 python-docx：{e}")

    document = docx.Document()

    def set_style_font(style):
        try:
            style.font.name = options.font_name
            style.font.size = Pt(11)
            rFonts = style.element.rPr.rFonts
            rFonts.set(qn('w:eastAsia'), options.font_name)
        except Exception:
            # 忽略字体设置失败
            pass

    set_style_font(document.styles['Normal'])

    document.add_heading(options.title, level=1)

    blocks = parse_structure(text)
    for b in blocks:
        if time.time() - start > options.timeout_seconds:
            raise TimeoutError("DOCX 转换超时")
        if b["type"] == "heading":
            lvl = b.get("level", 2)
            document.add_heading(b["text"], level=min(max(lvl, 1), 4))
        elif b["type"] == "paragraph":
            p = document.add_paragraph()
            run = p.add_run(b["text"])  # 保留换行由 Word 自行处理
            try:
                run.font.name = options.font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), options.font_name)
            except Exception:
                pass
        elif b["type"] == "list":
            # python-docx 无内置列表样式 API，此处用手工缩进与符号表示
            ordered = b.get("ordered")
            for idx, it in enumerate(b.get("items", []), start=1):
                bullet = f"{idx}. " if ordered else "• "
                p = document.add_paragraph(bullet + it)
                try:
                    p.paragraph_format.left_indent = Pt(12)
                except Exception:
                    pass

    document.save(output_path)
    return output_path


def convert_text_bytes(data: bytes, output_path: str, options: ConversionOptions) -> str:
    text = decode_text(data)
    if options.target_format.lower() == 'pdf':
        return convert_to_pdf(text, output_path, options)
    elif options.target_format.lower() == 'docx':
        return convert_to_docx(text, output_path, options)
    else:
        raise ValueError("不支持的目标格式，仅支持 pdf 或 docx")