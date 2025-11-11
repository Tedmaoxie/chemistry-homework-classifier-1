"""
分层DOCX输出生成器
使用 python-docx 生成结构化的分层化学作业报告（支持可编辑样式与目录）
"""

import logging
from dataclasses import dataclass
from typing import List, Optional, Dict, Any
from datetime import datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.difficulty_assessor.classifier import DifficultyLevel
from .pdf_generator import ClassifiedQuestion

logger = logging.getLogger(__name__)


@dataclass
class DOCXReportStyle:
    title: str = "高中化学分层作业"
    author: str = "chemistry-homework-classifier"


class ClassifiedDOCXGenerator:
    def __init__(self, style: Optional[DOCXReportStyle] = None):
        self.style = style or DOCXReportStyle()

    @staticmethod
    def _difficulty_name(level: DifficultyLevel) -> str:
        mapping = {
            DifficultyLevel.BEGINNER: "基础",
            DifficultyLevel.INTERMEDIATE: "中级",
            DifficultyLevel.ADVANCED: "高级",
        }
        return mapping.get(level, str(level.value))

    @staticmethod
    def _score(item: ClassifiedQuestion) -> float:
        try:
            return float(item.classification.detailed_analysis.get('total_score', 0))
        except Exception:
            return 0.0

    def _add_cover(self, doc: Document, doc_info: Dict[str, Any]) -> None:
        # 标题
        p = doc.add_paragraph(doc_info.get("title", self.style.title), style="Title")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 说明与副标题
        subtitle = doc_info.get("subtitle") or "化学作业分层与难度分析报告"
        doc.add_paragraph(subtitle, style="Intense Quote")

        sort_note = doc_info.get("sort_note")
        if sort_note:
            doc.add_paragraph(sort_note, style="Intense Quote")

        # 元信息表
        meta_rows = [
            ("生成时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
            ("作者", self.style.author),
            ("总题目数", str(doc_info.get("total_questions", "-"))),
        ]
        for k, v in doc_info.get("stats", {}).items():
            meta_rows.append((k, str(v)))

        table = doc.add_table(rows=len(meta_rows), cols=2)
        table.style = "Light Shading"
        for i, (k, v) in enumerate(meta_rows):
            table.cell(i, 0).text = str(k)
            table.cell(i, 1).text = str(v)

    def _insert_toc(self, doc: Document, heading_text: str = "目录") -> None:
        # 目录标题
        doc.add_paragraph(heading_text, style="Heading 1")
        # 插入 TOC 字段（需在Office中刷新字段以显示目录）
        p = doc.add_paragraph()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')

        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = 'TOC \\o "1-2" \\h \\z \\u'

        fld_char_separate = OxmlElement('w:fldChar')
        fld_char_separate.set(qn('w:fldCharType'), 'separate')

        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')

        r = OxmlElement('w:r')
        r.append(fld_char_begin)
        r.append(instr_text)
        r.append(fld_char_separate)
        r.append(fld_char_end)
        p._p.append(r)

    def _add_question_block(self, doc: Document, idx: int, item: ClassifiedQuestion) -> None:
        # 标题（进入目录）
        doc.add_paragraph(f"题目 {idx}", style="Heading 2")

        # 题干/选项
        for line in (item.question_text or '').splitlines():
            doc.add_paragraph(line, style="Normal")

        # 答案
        if item.answer_text:
            doc.add_paragraph("答案", style="Intense Quote")
            for line in (item.answer_text or '').splitlines():
                doc.add_paragraph(line, style="Normal")

        # 分类与分数
        conf_txt = f"分类级别: {self._difficulty_name(item.classification.level)} / 置信度: {item.classification.confidence:.2f}"
        doc.add_paragraph(conf_txt, style="Intense Quote")

        total_score = self._score(item)
        doc.add_paragraph(f"综合难度评分: {total_score:.1f}", style="Intense Quote")

        # 元数据（ID、标签等）
        if item.metadata:
            rows = []
            qid = item.metadata.get("question_id")
            if qid:
                rows.append(("题目ID", str(qid)))
            tags = item.metadata.get("knowledge_tags")
            if tags:
                rows.append(("知识点标签", ", ".join(map(str, tags))))
            for k, v in item.metadata.items():
                if k in ("question_id", "knowledge_tags"):
                    continue
                rows.append((str(k), str(v)))
            if rows:
                table = doc.add_table(rows=len(rows), cols=2)
                table.style = "Light Shading"
                for i, (k, v) in enumerate(rows):
                    table.cell(i, 0).text = k
                    table.cell(i, 1).text = v

    def generate_sorted_report(
        self,
        output_path: str,
        items: List[ClassifiedQuestion],
        metadata: Optional[Dict[str, Any]] = None,
    ) -> str:
        # 分组保持逻辑关联（如有 group_id），组间按最大分数降序，组内按分数降序
        def group_key(it: ClassifiedQuestion):
            gid = None
            if it.metadata:
                gid = it.metadata.get('group_id')
            return gid

        groups: Dict[Any, List[ClassifiedQuestion]] = {}
        for idx, it in enumerate(items):
            gid = group_key(it)
            key = gid if gid is not None else f"__single__{idx}"
            groups.setdefault(key, []).append(it)

        sorted_groups = sorted(
            groups.items(),
            key=lambda kv: max(self._score(x) for x in kv[1]),
            reverse=True,
        )
        sorted_items: List[ClassifiedQuestion] = []
        for _, g in sorted_groups:
            sorted_items.extend(sorted(g, key=lambda x: self._score(x), reverse=True))

        # 统计
        scores = [self._score(it) for it in sorted_items]
        total = len(sorted_items)
        stats = {
            "题目总数": total,
            "最高分": f"{max(scores) if scores else 0:.1f}",
            "最低分": f"{min(scores) if scores else 0:.1f}",
            "平均分": f"{(sum(scores)/total) if total else 0:.1f}",
        }

        # 文档生成
        doc = Document()
        self._add_cover(doc, {
            "title": self.style.title,
            "subtitle": (metadata or {}).get("subtitle"),
            "total_questions": total,
            "stats": stats,
            "sort_note": "排序依据：综合难度评分降序",
        })

        # 目录（需要在Word里“更新域”以展示）
        self._insert_toc(doc)

        # 正文标题
        doc.add_paragraph(f"题目列表（按综合难度评分降序，共{total}题）", style="Heading 1")

        # 题目块
        for idx, item in enumerate(sorted_items, start=1):
            self._add_question_block(doc, idx, item)

        doc.save(output_path)
        logger.info(f"按难度分数降序DOCX报告已生成: {output_path}")
        return output_path


def create_docx_sorted(output_path: str, items: List[ClassifiedQuestion], metadata: Optional[Dict[str, Any]] = None) -> str:
    gen = ClassifiedDOCXGenerator()
    return gen.generate_sorted_report(output_path, items, metadata)